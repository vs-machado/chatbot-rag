import io
import json
import uuid
from typing import Optional

import docx
from fastapi import HTTPException, UploadFile
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from sqlalchemy import desc
from sqlalchemy.orm import Session

from config import CHUNK_OVERLAP, CHUNK_SIZE
from database import SessionLocal
from models import Document
from schemas.document import DocumentCreate, DocumentSearchResult, ModelConfig
from services.embedding_service import generate_embedding, generate_embeddings
from services.upload_job_service import update_upload_job


def extract_text_from_stream(content_type: str | None, file_stream: io.BytesIO) -> str:
    """Extrai texto de um arquivo baseado no tipo MIME."""
    content = ""
    content_type = content_type or ""
    file_stream.seek(0)

    if content_type == "application/pdf":
        reader = PdfReader(file_stream)
        for page in reader.pages:
            content += page.extract_text() + "\n"

    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file_stream)
        for para in doc.paragraphs:
            content += para.text + "\n"

    elif content_type.startswith("text/"):
        content = file_stream.read().decode("utf-8")

    else:
        raise HTTPException(
            status_code=400, detail=f"Tipo de arquivo não suportado: {content_type}"
        )

    return content


def extract_text_from_file(file: UploadFile) -> str:
    """Extrai texto de um arquivo enviado via FastAPI."""
    return extract_text_from_stream(file.content_type, file.file)


def create_document(db: Session, doc_data: DocumentCreate) -> Document:
    """Cria um novo documento e gera seu embedding."""
    provider = None
    model = None
    api_key_google = None
    api_key_openai = None

    if doc_data.model_config_:
        provider = doc_data.model_config_.provider
        model = doc_data.model_config_.model
        if doc_data.model_config_.api_keys:
            api_key_google = doc_data.model_config_.api_keys.google_api_key
            api_key_openai = doc_data.model_config_.api_keys.openai_api_key

    embedding = generate_embedding(
        doc_data.content,
        provider=provider,
        model=model,
        google_api_key=api_key_google,
        openai_api_key=api_key_openai,
    )

    # Verifica dimensão do embedding.
    # Se mudar o modelo no meio do caminho sem migrar, vai dar erro no banco se dimensão for diferente
    # O ideal seria validar contra a coluna do banco, mas aqui assumimos que o user sabe o que faz
    # ou que estamos usando o padrão.

    db_doc = Document(
        title=doc_data.title,
        content=doc_data.content,
        embedding=embedding,
        metadata_=str(doc_data.metadata) if doc_data.metadata else None,
    )

    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    return db_doc


def process_and_create_documents_from_file(
    db: Session,
    file: UploadFile,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    model_config: Optional[ModelConfig] = None,
) -> list[Document]:
    """Processa um arquivo, divide em chunks e cria documentos."""
    text = extract_text_from_file(file)

    return create_documents_from_text(
        db=db,
        text=text,
        filename=file.filename or "arquivo",
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        model_config=model_config,
    )


def create_documents_from_text(
    db: Session,
    text: str,
    filename: str,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    model_config: Optional[ModelConfig] = None,
) -> list[Document]:
    """Cria documentos a partir de texto já extraído."""
    if not text.strip():
        raise HTTPException(status_code=400, detail="Arquivo vazio ou sem texto extraível")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )

    chunks = text_splitter.split_text(text)
    if not chunks:
        raise HTTPException(status_code=400, detail="Nenhum chunk válido foi gerado")

    created_docs = []

    # Extrai configs
    provider = model_config.provider if model_config else None
    model = model_config.model if model_config else None
    api_key_google = (
        model_config.api_keys.google_api_key if model_config and model_config.api_keys else None
    )
    api_key_openai = (
        model_config.api_keys.openai_api_key if model_config and model_config.api_keys else None
    )

    embeddings = generate_embeddings(
        chunks,
        provider=provider,
        model=model,
        google_api_key=api_key_google,
        openai_api_key=api_key_openai,
    )

    for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        doc = Document(
            title=f"{filename} - Parte {i + 1}",
            content=chunk,
            embedding=embedding,
            metadata_=json.dumps(
                {"source": filename, "chunk": i, "total_chunks": len(chunks)},
                ensure_ascii=True,
            ),
        )
        db.add(doc)
        created_docs.append(doc)

    db.commit()
    for doc in created_docs:
        db.refresh(doc)

    return created_docs


def process_and_create_documents_in_background(
    job_id: uuid.UUID,
    filename: str,
    content_type: str,
    file_bytes: bytes,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    model_config: Optional[ModelConfig] = None,
) -> None:
    """Processa documentos em background e atualiza progresso do job."""
    db = SessionLocal()
    try:
        update_upload_job(
            job_id,
            status="processing",
            message="Extraindo texto do arquivo.",
            progress_percentage=10,
        )
        text = extract_text_from_stream(content_type, io.BytesIO(file_bytes))
        if not text.strip():
            raise HTTPException(status_code=400, detail="Arquivo vazio ou sem texto extraível")

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        chunks = text_splitter.split_text(text)
        if not chunks:
            raise HTTPException(status_code=400, detail="Nenhum chunk válido foi gerado")

        update_upload_job(
            job_id,
            message="Gerando chunks do documento.",
            progress_percentage=25,
            total_chunks=len(chunks),
            processed_chunks=0,
        )

        provider = model_config.provider if model_config else None
        model = model_config.model if model_config else None
        api_key_google = (
            model_config.api_keys.google_api_key if model_config and model_config.api_keys else None
        )
        api_key_openai = (
            model_config.api_keys.openai_api_key if model_config and model_config.api_keys else None
        )

        update_upload_job(
            job_id,
            message="Gerando embeddings do documento.",
            progress_percentage=45,
        )
        embeddings = generate_embeddings(
            chunks,
            provider=provider,
            model=model,
            google_api_key=api_key_google,
            openai_api_key=api_key_openai,
        )

        created_docs = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            doc = Document(
                title=f"{filename} - Parte {i + 1}",
                content=chunk,
                embedding=embedding,
                metadata_=json.dumps(
                    {"source": filename, "chunk": i, "total_chunks": len(chunks)},
                    ensure_ascii=True,
                ),
            )
            db.add(doc)
            created_docs.append(doc)

        update_upload_job(
            job_id,
            message="Salvando documento na base de conhecimento.",
            progress_percentage=85,
            processed_chunks=len(chunks),
        )

        db.commit()
        for doc in created_docs:
            db.refresh(doc)

        update_upload_job(
            job_id,
            status="completed",
            message="Documento processado com sucesso.",
            progress_percentage=100,
            processed_chunks=len(chunks),
            documents_created=len(created_docs),
            document_ids=[doc.id for doc in created_docs],
        )
    except HTTPException as exc:
        db.rollback()
        update_upload_job(
            job_id,
            status="failed",
            message="Falha ao processar documento.",
            error=exc.detail,
        )
    except Exception as exc:
        db.rollback()
        update_upload_job(
            job_id,
            status="failed",
            message="Falha ao processar documento.",
            error=str(exc),
        )
    finally:
        db.close()


def list_documents(db: Session, skip: int = 0, limit: int = 100) -> list[Document]:
    """Lista documentos paginados."""
    return db.query(Document).order_by(desc(Document.created_at)).offset(skip).limit(limit).all()


def get_document(db: Session, document_id: uuid.UUID) -> Optional[Document]:
    """Busca documento por ID."""
    return db.query(Document).filter(Document.id == document_id).first()


def delete_document(db: Session, document_id: uuid.UUID) -> bool:
    """Deleta um documento."""
    doc = get_document(db, document_id)
    if not doc:
        return False
    db.delete(doc)
    db.commit()
    return True


def search_documents(
    db: Session, query: str, top_k: int = 5, model_config: Optional[ModelConfig] = None
) -> list[DocumentSearchResult]:
    """Busca documentos semanticamente similares."""
    # Extrai configs
    provider = model_config.provider if model_config else None
    model = model_config.model if model_config else None
    api_key_google = (
        model_config.api_keys.google_api_key if model_config and model_config.api_keys else None
    )
    api_key_openai = (
        model_config.api_keys.openai_api_key if model_config and model_config.api_keys else None
    )

    # Gera embedding da query
    query_embedding = generate_embedding(
        query,
        provider=provider,
        model=model,
        google_api_key=api_key_google,
        openai_api_key=api_key_openai,
        is_query=True,
    )

    # Busca usando pgvector (distância de cosseno <=> l2_distance se normalizado, mas pgvector tem operador específico)
    # A classe Document tem a coluna embedding.
    # Usamos order_by(Document.embedding.cosine_distance(query_embedding))

    results = (
        db.query(Document, Document.embedding.cosine_distance(query_embedding).label("distance"))
        .order_by(Document.embedding.cosine_distance(query_embedding))
        .limit(top_k)
        .all()
    )

    search_results = []
    for doc, distance in results:
        # Convertemos distância para similaridade (aproximado)
        # Cosine distance é 1 - cosine similarity.
        similarity = 1 - distance
        search_results.append(DocumentSearchResult(document=doc, score=similarity))

    return search_results
